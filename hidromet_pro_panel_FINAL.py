
import streamlit as st
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
from io import BytesIO
from docx import Document
from fpdf import FPDF
from datetime import datetime
import base64

# Configuración de página
st.set_page_config(page_title="HidroMet Pro Panel", layout="wide")

# --------- Autenticación ---------
if "authenticated" not in st.session_state:
    st.session_state.authenticated = False

def login():
    with st.form("login_form"):
        st.title("Iniciar sesión")
        usuario = st.text_input("Usuario", value="", max_chars=20)
        password = st.text_input("Contraseña", type="password")
        submitted = st.form_submit_button("Iniciar sesión")
        if submitted:
            if usuario == "admin" and password == "admin123":
                st.session_state.authenticated = True
                st.success(f"✅ Login exitoso. Bienvenido, {usuario}")
            else:
                st.error("❌ Usuario o contraseña incorrectos")

def logout():
    st.session_state.clear()

# --------- Funciones para exportar ---------
def generar_pdf(df):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.cell(200, 10, txt="Informe de Datos Hidrometeorológicos", ln=True, align='C')
    pdf.ln(10)
    for i in range(min(10, len(df))):
        row = df.iloc[i].to_string()
        for line in row.split("\n"):
            pdf.cell(200, 10, txt=line.strip(), ln=True)
        pdf.ln(5)
    pdf_output = BytesIO()
    pdf_bytes = pdf.output(dest='S').encode('latin1')
    pdf_output.write(pdf_bytes)
    pdf_output.seek(0)
    return pdf_output

def generar_word(df):
    doc = Document()
    doc.add_heading("Informe de Datos Hidrometeorológicos", 0)
    table = doc.add_table(rows=1, cols=len(df.columns))
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr_cells[i].text = str(col)
    for i in range(min(10, len(df))):
        row_cells = table.add_row().cells
        for j, val in enumerate(df.iloc[i]):
            row_cells[j].text = str(val)
    word_output = BytesIO()
    doc.save(word_output)
    word_output.seek(0)
    return word_output

# --------- Panel de administración ---------
def admin_panel():
    st.sidebar.title("⚙️ Panel de administración")
    if st.sidebar.button("Cerrar sesión"):
        logout()
        st.experimental_rerun()

    st.subheader("Sube tu archivo de datos (.csv)")
    uploaded_file = st.file_uploader("Carga un archivo", type=["csv"])

    if uploaded_file:
        df = pd.read_csv(uploaded_file)
        st.success("✅ Archivo cargado correctamente")

        st.markdown("## Vista previa de los datos")
        st.dataframe(df.head(10), use_container_width=True)

        st.markdown("## Gráficos interactivos")
        st.plotly_chart(px.line(df, x=df.columns[0], y=df.columns[1], title="Serie temporal"))
        st.plotly_chart(px.histogram(df, x=df.columns[1], title="Histograma de Variable"))
        st.plotly_chart(px.scatter(df, x=df.columns[0], y=df.columns[1], title="Dispersión"))
        st.plotly_chart(px.box(df, y=df.columns[1], title="Boxplot de Variable"))

        st.markdown("## Análisis Estadístico")
        st.write(df.describe())

        st.markdown("## Descarga de informes")
        col1, col2 = st.columns(2)
        with col1:
            pdf_file = generar_pdf(df)
            st.download_button("📄 Descargar PDF", data=pdf_file, file_name="informe.pdf", mime="application/pdf")
        with col2:
            word_file = generar_word(df)
            st.download_button("📝 Descargar Word", data=word_file, file_name="informe.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    else:
        st.warning("Por favor cargue un archivo para acceder a las herramientas")

# --------- MAIN ---------
def main():
    if st.session_state.authenticated:
        admin_panel()
    else:
        login()

if __name__ == "__main__":
    main()
