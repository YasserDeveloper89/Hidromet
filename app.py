import streamlit as st
import pandas as pd
import numpy as np
import plotly.express as px
import plotly.graph_objects as go
from io import BytesIO
from docx import Document
from fpdf import FPDF
from datetime import datetime

# Simulamos una base de datos de usuarios
USUARIOS = {
    "admin": {"password": "admin123", "role": "admin"},
    "tecnico": {"password": "tecnico123", "role": "tecnico"},
}

def login():
    st.title("🔐 Inicio de Sesión Corporativo")

    if "autenticado" not in st.session_state:
        st.session_state.autenticado = False
        st.session_state.usuario = None
        st.session_state.rol = None

    if not st.session_state.autenticado:
        with st.form("login_form"):
            username = st.text_input("Usuario")
            password = st.text_input("Contraseña", type="password")
            login_btn = st.form_submit_button("Iniciar sesión")

            if login_btn:
                if username in USUARIOS and USUARIOS[username]["password"] == password:
                    st.session_state.autenticado = True
                    st.session_state.usuario = username
                    st.session_state.rol = USUARIOS[username]["role"]
                    st.success(f"✅ Login exitoso. Bienvenido, {username}")
                else:
                    st.error("❌ Credenciales inválidas. Intenta de nuevo.")
    else:
        st.success(f"✅ Ya has iniciado sesión como: {st.session_state.usuario}")

def logout():
    st.session_state.autenticado = False
    st.session_state.usuario = None
    st.session_state.rol = None
    st.experimental_rerun()

# ---------------------- Panel Admin ----------------------

def panel_admin():
    st.sidebar.title("Panel de Administrador")
    st.sidebar.button("Cerrar sesión", on_click=logout)

    st.title("📊 Herramientas de Gestión Avanzada")

    uploaded_file = st.file_uploader("📤 Cargar archivo CSV con datos de sensores", type=["csv"])
    if uploaded_file:
        df = pd.read_csv(uploaded_file)

        st.subheader("🔎 Vista previa de datos")
        st.dataframe(df.head())

        st.subheader("📌 Estadísticas descriptivas")
        st.write(df.describe())

        st.subheader("📈 Gráficos interactivos")
        try:
            numeric_cols = df.select_dtypes(include=np.number).columns
            col1 = st.selectbox("Eje X", numeric_cols)
            col2 = st.selectbox("Eje Y", numeric_cols)
            fig = px.scatter(df, x=col1, y=col2, title="Dispersión de datos")
            st.plotly_chart(fig)
        except:
            st.warning("No se pueden generar gráficos con columnas seleccionadas.")

        st.subheader("🌐 Mapa de correlación")
        try:
            fig_corr = px.imshow(df[numeric_cols].corr(), text_auto=True, aspect="auto")
            st.plotly_chart(fig_corr)
        except:
            st.warning("No se pudo generar el mapa de correlación.")

        st.subheader("📤 Exportar informes")
        if st.button("Exportar como PDF"):
            try:
                pdf = FPDF()
                pdf.add_page()
                pdf.set_font("Arial", size=12)
                pdf.cell(200, 10, txt="Informe de Datos", ln=1, align="C")

                for i in range(len(df)):
                    row = ', '.join([str(x) for x in df.iloc[i]])
                    pdf.cell(200, 10, txt=row, ln=1, align="L")
                pdf_output = BytesIO()
                pdf.output(pdf_output)
                st.download_button("📄 Descargar PDF", data=pdf_output.getvalue(), file_name="informe.pdf")
            except Exception as e:
                st.error(f"Error al generar PDF: {e}")

        if st.button("Exportar como Word"):
            try:
                doc = Document()
                doc.add_heading("Informe de Datos", 0)
                table = doc.add_table(rows=1, cols=len(df.columns))
                hdr_cells = table.rows[0].cells
                for i, col in enumerate(df.columns):
                    hdr_cells[i].text = str(col)
                for index, row in df.iterrows():
                    row_cells = table.add_row().cells
                    for i, item in enumerate(row):
                        row_cells[i].text = str(item)
                doc_output = BytesIO()
                doc.save(doc_output)
                st.download_button("📝 Descargar Word", data=doc_output.getvalue(), file_name="informe.docx")
            except Exception as e:
                st.error(f"Error al generar Word: {e}")

        # 🎯 Funcionalidades extra (10 más)
        st.subheader("📡 Simulación de conexión en tiempo real")
        if st.button("🔄 Simular lectura de sensores"):
            data = {"Sensor1": np.random.random(), "Sensor2": np.random.random()}
            st.json(data)

        st.subheader("📌 Filtros personalizados")
        col = st.selectbox("Selecciona una columna a filtrar", df.columns)
        unique_vals = df[col].unique()
        selected_val = st.selectbox("Valor", unique_vals)
        st.dataframe(df[df[col] == selected_val])

        st.subheader("📍 Histograma de variables")
        hist_col = st.selectbox("Columna para histograma", numeric_cols)
        st.plotly_chart(px.histogram(df, x=hist_col))

        st.subheader("📊 Boxplot")
        st.plotly_chart(px.box(df, y=hist_col))

        st.subheader("📌 Detalle de outliers")
        q1 = df[hist_col].quantile(0.25)
        q3 = df[hist_col].quantile(0.75)
        iqr = q3 - q1
        outliers = df[(df[hist_col] < q1 - 1.5 * iqr) | (df[hist_col] > q3 + 1.5 * iqr)]
        st.write(outliers)

        st.subheader("🧮 Tendencia de series temporales")
        if "fecha" in df.columns:
            try:
                df["fecha"] = pd.to_datetime(df["fecha"])
                df.set_index("fecha", inplace=True)
                st.line_chart(df[numeric_cols])
            except:
                st.warning("No se pudo analizar la serie temporal")

        st.subheader("🔀 Análisis de componentes principales (PCA)")
        try:
            from sklearn.decomposition import PCA
            pca = PCA(n_components=2)
            components = pca.fit_transform(df[numeric_cols].dropna())
            fig_pca = px.scatter(x=components[:, 0], y=components[:, 1], labels={"x": "PC1", "y": "PC2"})
            st.plotly_chart(fig_pca)
        except Exception as e:
            st.warning("No se pudo calcular PCA")

        st.subheader("📂 Análisis de agrupamiento (K-Means)")
        try:
            from sklearn.cluster import KMeans
            km = KMeans(n_clusters=3)
            df_clean = df[numeric_cols].dropna()
            df['cluster'] = km.fit_predict(df_clean)
            st.plotly_chart(px.scatter(df_clean, x=numeric_cols[0], y=numeric_cols[1], color=df['cluster'].astype(str)))
        except Exception as e:
            st.warning("No se pudo ejecutar KMeans")

        st.subheader("📡 Conexión a dispositivos IoT")
        st.info("Simulación: Dispositivo conectado con éxito. Última actualización hace 5 segundos.")

        st.subheader("🧭 Indicadores dinámicos")
        col1, col2 = st.columns(2)
        col1.metric("Temperatura", f"{np.random.randint(20, 30)} °C", "+1")
        col2.metric("Humedad", f"{np.random.randint(40, 60)} %", "-2")

    else:
        st.warning("⚠️ Por favor cargue un archivo para acceder a las herramientas.")

def panel_tecnico():
    st.title("🔧 Panel de Técnico")
    st.write("Bienvenido al panel de técnico. Aquí irán las funciones correspondientes.")
    st.sidebar.button("Cerrar sesión", on_click=logout)

def main():
    if "autenticado" in st.session_state and st.session_state.autenticado:
        if st.session_state.rol == "admin":
            panel_admin()
        elif st.session_state.rol == "tecnico":
            panel_tecnico()
    else:
        login()

if __name__ == "__main__":
    main()
                                 
