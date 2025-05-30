
import streamlit as st
import pandas as pd
import numpy as np
import plotly.express as px
import plotly.graph_objects as go
from io import BytesIO
from docx import Document
from fpdf import FPDF
from datetime import datetime
import base64

# ------------------------ FUNCIONES ------------------------
def generar_pdf(df):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.cell(200, 10, txt="Informe de Datos", ln=True, align='C')
    pdf.ln(10)
    for col in df.columns:
        pdf.cell(200, 10, txt=f"{col}: {df[col].iloc[0]}", ln=True)
    buffer = BytesIO()
    pdf.output(buffer)
    buffer.seek(0)
    return buffer.read()

def generar_word(df):
    doc = Document()
    doc.add_heading('Informe de Datos', 0)
    table = doc.add_table(rows=1, cols=len(df.columns))
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr_cells[i].text = col
    row_cells = table.add_row().cells
    for i, val in enumerate(df.iloc[0]):
        row_cells[i].text = str(val)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()

def logout():
    st.session_state.logged_in = False
    st.session_state.username = None

def login():
    st.title("Panel de Inicio de Sesión")
    username = st.text_input("Usuario")
    password = st.text_input("Contraseña", type="password")
    if st.button("Acceder"):
        if username == "admin" and password == "admin123":
            st.session_state.logged_in = True
            st.session_state.username = username
            st.rerun()
        else:
            st.error("Credenciales inválidas. Inténtelo de nuevo.")

# ------------------------ PANEL ADMIN ------------------------
def admin_panel():
    st.sidebar.title("Panel de Control - Admin")
    if st.sidebar.button("Cerrar Sesión"):
        logout()
        st.rerun()

    st.title("Panel de Administración - Herramientas Avanzadas")
    uploaded_file = st.file_uploader("Cargue su archivo CSV", type="csv")
    if uploaded_file:
        df = pd.read_csv(uploaded_file)
        st.success("Datos cargados exitosamente.")
        
        # Herramienta 1: Vista previa de datos
        st.subheader("Vista previa de los datos")
        st.dataframe(df.head())

        # Herramienta 2: Estadísticas descriptivas
        st.subheader("Estadísticas descriptivas")
        st.write(df.describe())

        # Herramienta 3: Gráfico de correlación (solo columnas numéricas)
        st.subheader("Matriz de correlación")
        numeric_df = df.select_dtypes(include='number')
        if not numeric_df.empty:
            st.write(px.imshow(numeric_df.corr(), text_auto=True, title="Correlaciones"))

        # Herramienta 4: Histograma interactivo
        st.subheader("Histograma")
        column = st.selectbox("Seleccione columna", df.select_dtypes(include='number').columns)
        st.plotly_chart(px.histogram(df, x=column))

        # Herramienta 5: Serie temporal si hay fecha
        st.subheader("Serie Temporal")
        date_cols = df.select_dtypes(include='object').columns[df.select_dtypes(include='object').apply(lambda x: pd.to_datetime(x, errors='coerce').notnull().all())]
        if not date_cols.empty:
            date_col = date_cols[0]
            df[date_col] = pd.to_datetime(df[date_col])
            df.set_index(date_col, inplace=True)
            st.line_chart(df.select_dtypes(include='number'))

        # Herramienta 6: Filtrado dinámico
        st.subheader("Filtrado dinámico")
        filtro_col = st.selectbox("Filtrar por", df.columns)
        valores = st.multiselect("Valores", df[filtro_col].unique())
        if valores:
            st.write(df[df[filtro_col].isin(valores)])

        # Herramienta 7: Descargar PDF
        st.subheader("Exportar informe")
        if st.button("Descargar informe en PDF"):
            pdf_data = generar_pdf(df)
            b64 = base64.b64encode(pdf_data).decode()
            href = f'<a href="data:application/pdf;base64,{b64}" download="informe.pdf">Descargar PDF</a>'
            st.markdown(href, unsafe_allow_html=True)

        # Herramienta 8: Descargar Word
        if st.button("Descargar informe en Word"):
            word_data = generar_word(df)
            b64 = base64.b64encode(word_data).decode()
            href = f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="informe.docx">Descargar Word</a>'
            st.markdown(href, unsafe_allow_html=True)
    else:
        st.info("Por favor cargue un archivo para acceder a las herramientas.")

# ------------------------ MAIN ------------------------
def main():
    if "logged_in" not in st.session_state:
        st.session_state.logged_in = False
    if st.session_state.logged_in:
        admin_panel()
    else:
        login()

if __name__ == "__main__":
    main()
