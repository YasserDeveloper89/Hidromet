import streamlit as st
import pandas as pd
import plotly.express as px
from fpdf import FPDF
from io import BytesIO
from docx import Document
from sklearn.linear_model import LinearRegression
import numpy as np

# ----------------- Autenticación -----------------
USUARIOS = {
    "admin": "YZ1BKzgHIK7P7ZrB",
    "tecnico": "tecnico123"
}

# ----------------- Login -----------------
def login():
    st.title("💧 Hydromet - Centro de control de datos ambientales")
    usuario = st.text_input("Usuario")
    contraseña = st.text_input("Contraseña", type="password")
    if st.button("Iniciar sesión"):
        if usuario in USUARIOS and USUARIOS[usuario] == contraseña:
            st.session_state.autenticado = True
            st.session_state.usuario = usuario
            st.success(f"✅ Login exitoso. Bienvenido, {usuario}")
            st.rerun()
        else:
            st.error("❌ Usuario o contraseña incorrectos")

# ----------------- Logout -----------------
def logout():
    st.session_state.autenticado = False
    st.session_state.usuario = ""
    st.session_state.df_cargado = None
    st.rerun()

# ----------------- Exportar a PDF -----------------
def generar_pdf(df):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.cell(200, 10, txt="Reporte de Datos", ln=True, align="C")
    pdf.ln()
    col_width = pdf.w / (len(df.columns) + 1)
    for col in df.columns:
        pdf.cell(col_width, 10, str(col), border=1)
    pdf.ln()
    for row in df.itertuples(index=False):
        for item in row:
            pdf.cell(col_width, 10, str(item), border=1)
        pdf.ln()
    return pdf.output(dest='S').encode('latin-1')

# ----------------- Exportar a Word -----------------
def generar_word(df):
    doc = Document()
    doc.add_heading("Reporte de Datos", 0)
    table = doc.add_table(rows=1, cols=len(df.columns))
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr_cells[i].text = col
    for index, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, col_name in enumerate(df.columns):
            row_cells[i].text = str(row[col_name])
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()

# ----------------- Predicción ML -----------------
def modelo_prediccion(df):
    st.subheader("🤖 Predicción Automática")
    numeric_df = df.select_dtypes(include='number')
    if numeric_df.shape[1] < 2:
        st.warning("Se necesitan al menos dos columnas numéricas para predecir.")
        return

    cols = numeric_df.columns.tolist()
    x_col = st.selectbox("Variable independiente (X)", cols, key="pred_x")
    y_col = st.selectbox("Variable dependiente (Y)", [c for c in cols if c != x_col], key="pred_y")

    X = df[[x_col]].dropna()
    y = df[y_col].loc[X.index]

    modelo = LinearRegression()
    modelo.fit(X, y)
    prediccion = modelo.predict(X)

    fig = px.scatter(df, x=x_col, y=y_col, title="Predicción vs Realidad")
    fig.add_scatter(x=X[x_col], y=prediccion, mode='lines', name='Predicción')
    st.plotly_chart(fig)

    nuevo_valor = st.number_input(f"Ingresar nuevo valor para {x_col}")
    if st.button("Predecir"):
        resultado = modelo.predict([[nuevo_valor]])[0]
        st.success(f"✅ Predicción: {resultado:.2f}")

# ----------------- Análisis estadístico -----------------
def analisis_estadistico(df):
    st.subheader("📊 Análisis Estadístico")
    st.write("Resumen estadístico:")
    st.dataframe(df.describe())

    st.write("🔍 Outliers detectados:")
    for col in df.select_dtypes(include='number').columns:
        q1 = df[col].quantile(0.25)
        q3 = df[col].quantile(0.75)
        iqr = q3 - q1
        outliers = df[(df[col] < q1 - 1.5 * iqr) | (df[col] > q3 + 1.5 * iqr)]
        if not outliers.empty:
            st.warning(f"Columna '{col}' tiene {len(outliers)} outliers.")

# ----------------- PDF analítico -----------------
def generar_pdf_analitico(df):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.cell(200, 10, txt="Reporte Analítico", ln=True, align="C")
    pdf.ln()
    resumen = df.describe().round(2)
    for col in resumen.columns:
        pdf.cell(200, 10, txt=f"{col}:", ln=True)
        for idx in resumen.index:
            pdf.cell(200, 10, txt=f"{idx}: {resumen[col]}", ln=True)
        pdf.ln()
    return pdf.output(dest='S').encode('latin-1')

# ----------------- Cargar CSV -----------------
def cargar_datos():
    uploaded_file = st.file_uploader("📁 Cargar archivo CSV", type=["csv"])
    if uploaded_file:
        try:
            df = pd.read_csv(uploaded_file)
            if 'fecha' in df.columns:
                df['fecha'] = pd.to_datetime(df['fecha'])
                df.set_index('fecha', inplace=True)
                st.success("CSV cargado y columna 'fecha' establecida como índice.")
            st.session_state.df_cargado = df
        except Exception as e:
            st.error(f"Error al cargar el archivo: {e}")
            st.session_state.df_cargado = None

# ----------------- Panel del Administrador -----------------
def admin_panel():
    st.title("🛠️ Panel de Administración")
    st.write(f"Bienvenido, {st.session_state.usuario}")
    cargar_datos()

    df = st.session_state.get('df_cargado')
    if df is not None:
        st.subheader("📄 Vista Previa de los Datos")
        st.dataframe(df)

        st.subheader("📈 Serie de Tiempo")
        st.line_chart(df)

        numeric_df = df.select_dtypes(include='number')
        if not numeric_df.empty:
            st.subheader("📊 Correlación")
            fig = px.imshow(numeric_df.corr(), text_auto=True, title="Matriz de Correlación")
            st.plotly_chart(fig)

            st.subheader("📌 Dispersión")
            cols = numeric_df.columns.tolist()
            x = st.selectbox("Eje X", cols)
            y = st.selectbox("Eje Y", cols)
            if x and y:
                fig_scatter = px.scatter(df, x=x, y=y, title=f"{x} vs {y}")
                st.plotly_chart(fig_scatter)

            st.subheader("📌 Histograma")
            col_hist = st.selectbox("Selecciona columna", cols)
            fig_hist = px.histogram(df, x=col_hist, marginal="rug", title=f"Histograma de {col_hist}")
            st.plotly_chart(fig_hist)

        modelo_prediccion(df)
        analisis_estadistico(df)

        st.subheader("📄 Exportar Datos")
        st.download_button("📄 Descargar PDF", data=generar_pdf(df), file_name="reporte.pdf", mime="application/pdf")
        st.download_button("📝 Descargar Word", data=generar_word(df), file_name="reporte.docx",
                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

        st.subheader("📅 Exportar Reporte Inteligente")
        st.download_button("📁 Descargar Reporte Avanzado (PDF)",
                           data=generar_pdf_analitico(df),
                           file_name="reporte_analitico.pdf",
                           mime="application/pdf")

    if st.button("Cerrar sesión"):
        logout()

# ----------------- Panel del Técnico -----------------
def tecnico_panel():
    st.title("🔧 Panel Técnico")
    st.write(f"Bienvenido, {st.session_state.usuario}")
    cargar_datos()

    df = st.session_state.get('df_cargado')
    if df is not None:
        st.subheader("📄 Vista Previa de los Datos")
        st.dataframe(df)

        st.subheader("📈 Serie de Tiempo")
        st.line_chart(df)

        numeric_df = df.select_dtypes(include='number')
        if not numeric_df.empty:
            st.subheader("📊 Correlación")
            fig = px.imshow(numeric_df.corr(), text_auto=True, title="Matriz de Correlación")
            st.plotly_chart(fig)

        st.subheader("📄 Exportar Datos")
        st.download_button("📄 Descargar PDF", data=generar_pdf(df), file_name="reporte.pdf", mime="application/pdf")
        st.download_button("📝 Descargar Word", data=generar_word(df), file_name="reporte.docx",
                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    else:
        st.info("¡Comienza tu análisis! Sube un archivo CSV y visualiza los datos de forma organizada")

    if st.button("Cerrar sesión"):
        logout()

# ----------------- Inicialización -----------------
if 'autenticado' not in st.session_state:
    st.session_state.autenticado = False
if 'usuario' not in st.session_state:
    st.session_state.usuario = ""
if 'df_cargado' not in st.session_state:
    st.session_state.df_cargado = None

# ----------------- Main App -----------------
def main():
    if st.session_state.autenticado:
        if st.session_state.usuario == "admin":
            admin_panel()
        elif st.session_state.usuario == "tecnico":
            tecnico_panel()
    else:
        login()

main()

    
