
import streamlit as st
import pandas as pd
import plotly.express as px
from fpdf import FPDF
from io import BytesIO
from docx import Document

# ----------------- Estilos -----------------
st.markdown(
    """
    <style>
        html, body, [class*="css"] {
            font-family: 'Segoe UI', sans-serif;
            background-color: #1e1e1e;
            color: #f0e68c;
        }
        .stButton>button {
            background-color: #333333;
            color: #f0e68c;
            border-radius: 8px;
            padding: 0.5em 1em;
            border: none;
        }
        .stTextInput>div>div>input {
            background-color: #2a2a2a;
            color: #f0e68c;
            border-radius: 5px;
        }
    </style>
    """,
    unsafe_allow_html=True
)

# ----------------- Autenticación -----------------
USUARIOS = {
    "admin": "YZ1BKzgHIK7P7ZrB",
    "tecnico": "tecnico123"
}

# ----------------- Exportar a PDF -----------------
def generar_pdf(df):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.cell(200, 10, txt="Reporte de Datos", ln=True, align="C")
    pdf.ln()
    col_width = pdf.w / (len(df.columns) + 1)
    for col in df.columns:
        pdf.cell(col_width, 10, str(col), border=1)
    pdf.ln()
    for row in df.itertuples(index=False):
        for item in row:
            pdf.cell(col_width, 10, str(item), border=1)
        pdf.ln()
    return pdf.output(dest='S').encode('latin-1')

# ----------------- Exportar a Word -----------------
def generar_word(df):
    doc = Document()
    doc.add_heading("Reporte de Datos", 0)
    table = doc.add_table(rows=1, cols=len(df.columns))
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr_cells[i].text = col
    for index, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, col_name in enumerate(df.columns):
            row_cells[i].text = str(row[col_name])
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()

# ----------------- Login -----------------
def login():
    st.title("🔐 Hydromet - Ingreso al Sistema")
    usuario = st.text_input("Usuario")
    contraseña = st.text_input("Contraseña", type="password")
    if st.button("Iniciar sesión"):
        if usuario in USUARIOS and USUARIOS[usuario] == contraseña:
            st.session_state.autenticado = True
            st.session_state.usuario = usuario
        else:
            st.error("❌ Usuario o contraseña incorrectos")

# ----------------- Logout -----------------
def logout():
    st.session_state.autenticado = False
    st.session_state.usuario = ""
    st.session_state.df_cargado = None

# ----------------- Cargar CSV -----------------
def cargar_datos():
    uploaded_file = st.file_uploader("📁 Cargar archivo CSV", type=["csv"])
    if uploaded_file:
        try:
            df = pd.read_csv(uploaded_file)
            if 'fecha' in df.columns:
                df['fecha'] = pd.to_datetime(df['fecha'])
                df.set_index('fecha', inplace=True)
                st.success("✅ CSV cargado con 'fecha' como índice.")
            st.session_state.df_cargado = df
        except Exception as e:
            st.error(f"Error al cargar el archivo: {e}")
            st.session_state.df_cargado = None

# ----------------- Panel Técnico -----------------
def tecnico_panel():
    st.title("🔧 Panel Técnico")
    cargar_datos()
    df = st.session_state.get('df_cargado')
    if df is not None:
        st.subheader("📄 Vista Previa")
        st.dataframe(df)
        st.metric("Total de registros", len(df))
        st.subheader("📈 Serie de Tiempo")
        st.line_chart(df)
        numeric_df = df.select_dtypes(include='number')
        if not numeric_df.empty:
            st.subheader("📊 Matriz de Correlación")
            fig = px.imshow(numeric_df.corr(), text_auto=True)
            st.plotly_chart(fig)
        st.subheader("📤 Exportar")
        st.download_button("📄 PDF", data=generar_pdf(df), file_name="reporte.pdf", mime="application/pdf")
        st.download_button("📝 Word", data=generar_word(df), file_name="reporte.docx",
                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    if st.button("Cerrar sesión"):
        logout()

# ----------------- Panel Administrador -----------------
def admin_panel():
    st.title("🛠️ Panel de Administración")
    cargar_datos()
    df = st.session_state.get('df_cargado')
    if df is not None:
        st.subheader("📄 Vista Previa")
        st.dataframe(df)
        st.metric("Total de registros", len(df))
        st.subheader("📈 Serie de Tiempo")
        st.line_chart(df)
        numeric_df = df.select_dtypes(include='number')
        if not numeric_df.empty:
            st.subheader("📊 Matriz de Correlación")
            fig = px.imshow(numeric_df.corr(), text_auto=True)
            st.plotly_chart(fig)
            st.subheader("📌 Histograma")
            col_hist = st.selectbox("Columna", numeric_df.columns)
            fig_hist = px.histogram(df, x=col_hist)
            st.plotly_chart(fig_hist)
            st.subheader("📌 Boxplot")
            fig_box = px.box(df, y=col_hist)
            st.plotly_chart(fig_box)
            st.subheader("🧮 Nulos por columna")
            st.write(df.isnull().sum())
            st.subheader("📅 Distribución semanal")
            st.line_chart(df.resample('W').mean(numeric_only=True))
            st.subheader("📈 Estadísticas descriptivas")
            st.write(numeric_df.describe())
        st.subheader("📤 Exportar")
        st.download_button("📄 PDF", data=generar_pdf(df), file_name="reporte.pdf", mime="application/pdf")
        st.download_button("📝 Word", data=generar_word(df), file_name="reporte.docx",
                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    if st.button("Cerrar sesión"):
        logout()

# ----------------- App Principal -----------------
def main():
    if 'autenticado' not in st.session_state:
        st.session_state.autenticado = False
    if 'usuario' not in st.session_state:
        st.session_state.usuario = ""
    if 'df_cargado' not in st.session_state:
        st.session_state.df_cargado = None

    if st.session_state.autenticado:
        if st.session_state.usuario == "admin":
            admin_panel()
        elif st.session_state.usuario == "tecnico":
            tecnico_panel()
    else:
        login()

main()
